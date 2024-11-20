# helpers.py

import csv
from docx import Document
from docx.shared import Inches

class MediaPreparationHelper:
    # Conversion factors for units
    conversion_factors = {
        'mg': 1e-3,      # mg to grams
        'ug': 1e-6,      # ug to grams
        'g': 1,          # grams to grams
        'mg/mL': 1e3,    # mg/mL to μg/mL
        'ug/mL': 1,
        'ng/mL': 1e-3,   # ng/mL to μg/mL
        'ng/μL': 1,      # ng/μL to ng/μL (no conversion needed)
        'M': 1,          # M to M
        'mM': 1e-3,      # mM to M
        'uM': 1e-6,      # uM to M
        'nM': 1e-9,      # nM to M
        'X': 1,          # For components like ITS at 1X concentration
    }

    def __init__(self, components_stock, final_volume_ml, recipe_file):
        self.components_stock = components_stock
        self.final_volume_ml = final_volume_ml
        self.recipe_file = recipe_file

        # Initialize base media and serum with default values
        self.base_media = {
            'name': 'HEPES-buffered DMEM/F12',
            'type': 'Base Media',
        }
        self.serum = {
            'name': 'Fetal Bovine Serum (FBS)',
            'type': 'Serum',
            'percentage': 10,  # Default serum percentage
        }

        # Now call read_recipe after initializing base_media and serum
        self.recipe_data = self.read_recipe(recipe_file)

    def read_recipe(self, recipe_file):
        recipe = []

        with open(recipe_file, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            rows = list(reader)

            if len(rows) < 2:
                raise ValueError("CSV file must have at least two rows for base media and serum.")

            # First row: Base Media
            base_media_row = rows[0]
            if base_media_row['Type'] != 'Base Media':
                raise ValueError("First row must be the Base Media.")

            self.base_media['name'] = base_media_row['Name']

            # Second row: Serum
            serum_row = rows[1]
            if serum_row['Type'] != 'Serum':
                raise ValueError("Second row must be the Serum.")

            self.serum['name'] = serum_row['Name']
            if serum_row['Dilution Factor']:
                try:
                    self.serum['percentage'] = float(serum_row['Dilution Factor'])
                except ValueError:
                    raise ValueError("Invalid serum percentage provided.")

            # Remaining rows: Additives
            additive_rows = rows[2:]
            for row in additive_rows:
                if row['Type'] != 'Additive':
                    continue  # Skip if Type is not Additive
                component = {
                    'name': row['Name'],
                }
                if row['Dilution Factor']:
                    component['dilution_factor'] = float(row['Dilution Factor'])
                elif row['Desired Concentration'] and row['Desired Unit']:
                    component['desired_concentration'] = float(row['Desired Concentration'])
                    component['desired_unit'] = row['Desired Unit']
                else:
                    continue  # Skip if neither concentration nor dilution factor is provided
                recipe.append(component)

        return recipe

    def find_component_stock(self, name):
        for comp in self.components_stock:
            if comp['name'] == name:
                return comp
        return None

    def calculate_volume(self, component_stock, desired_concentration_info):
        final_volume_ml = self.final_volume_ml

        if 'dilution_factor' in desired_concentration_info:
            # Volume based on dilution factor
            volume_needed_ml = final_volume_ml / desired_concentration_info['dilution_factor']
            volume_needed_ul = volume_needed_ml * 1000
            return volume_needed_ul
        else:
            # Get stock concentration
            stock_concentration = component_stock.get('stock_concentration', component_stock.get('desired_stock_concentration'))
            stock_unit = component_stock.get('stock_unit')
            if stock_concentration is None or stock_unit is None:
                raise ValueError(f"No stock concentration found for component {component_stock['name']}")

            desired_concentration = desired_concentration_info['desired_concentration']
            desired_unit = desired_concentration_info['desired_unit']

            # Handle working solutions
            if 'working_solution_concentration' in component_stock:
                working_concentration = component_stock['working_solution_concentration']
                working_unit = component_stock['working_solution_unit']

                # Calculate dilution factor to make working solution
                dilution_factor = (stock_concentration * self.conversion_factors[stock_unit]) / \
                                  (working_concentration * self.conversion_factors[working_unit])

                component_stock['working_solution_dilution_factor'] = dilution_factor
                stock_concentration = working_concentration
                stock_unit = working_unit

            # Convert concentrations to common units
            if stock_unit not in self.conversion_factors or desired_unit not in self.conversion_factors:
                raise ValueError(f"Unit conversion not defined for units {stock_unit} or {desired_unit}")

            stock_conc_common = stock_concentration * self.conversion_factors[stock_unit]
            desired_conc_common = desired_concentration * self.conversion_factors[desired_unit]

            # Calculate volume in uL
            volume_ul = (desired_conc_common * final_volume_ml * 1000) / stock_conc_common

            # Adjust stock concentration if volume is less than 5 uL
            if volume_ul < 5:
                # Calculate new stock concentration to make volume_ul = 5 uL
                volume_ul = 5  # Set minimum volume to 5 uL
                stock_conc_common = (desired_conc_common * final_volume_ml * 1000) / volume_ul
                # Convert back to original stock unit
                new_stock_concentration = stock_conc_common / self.conversion_factors[stock_unit]
                # Update stock concentration in component_stock
                component_stock['adjusted_stock_concentration'] = new_stock_concentration
                component_stock['stock_concentration'] = new_stock_concentration
                # Recalculate volume_ul with the new stock concentration
                volume_ul = (desired_conc_common * final_volume_ml * 1000) / stock_conc_common

            return volume_ul

    def calculate_stock_solutions(self):
        stock_preparations = []
        for comp in self.components_stock:
            if 'initial_weight' in comp and 'desired_stock_concentration' in comp:
                initial_weight = comp['initial_weight']  # Amount
                initial_weight_unit = comp['initial_weight_unit']  # Unit
                molecular_weight = comp.get('molecular_weight')  # g/mol
                desired_concentration = comp.get('adjusted_stock_concentration', comp['desired_stock_concentration'])
                stock_unit = comp['stock_unit']
                solvent = comp.get('solvent', 'Appropriate solvent')

                # Convert initial weight to grams
                initial_weight_g = initial_weight * self.conversion_factors[initial_weight_unit]

                # Calculate volume_ml based on the adjusted concentration
                if stock_unit in ['mg/mL', 'μg/mL', 'ug/mL', 'ng/μL']:
                    # For mass-based concentrations
                    if stock_unit == 'mg/mL':
                        concentration_mg_per_ml = desired_concentration
                    elif stock_unit in ['μg/mL', 'ug/mL']:
                        concentration_mg_per_ml = desired_concentration / 1000  # μg/mL to mg/mL
                    elif stock_unit == 'ng/μL':
                        concentration_mg_per_ml = desired_concentration / 1e6  # ng/μL to mg/mL

                    mass_mg = initial_weight_g * 1000  # Convert g to mg
                    volume_ml = mass_mg / concentration_mg_per_ml  # in mL

                elif stock_unit in ['M', 'mM', 'μM', 'uM']:
                    # For molar concentrations
                    if molecular_weight is None:
                        raise ValueError(f"Molecular weight is required for component {comp['name']}")

                    desired_concentration_M = desired_concentration * self.conversion_factors[stock_unit]

                    moles = initial_weight_g / molecular_weight  # in mol
                    volume_L = moles / desired_concentration_M  # in L
                    volume_ml = volume_L * 1000  # in mL

                else:
                    volume_ml = None

                # Adjust initial weight down if volume exceeds 15 mL
                if volume_ml and volume_ml > 15:
                    volume_ml = 15  # Set volume to 15 mL
                    if stock_unit in ['mg/mL', 'μg/mL', 'ug/mL', 'ng/μL']:
                        # Adjust initial weight accordingly
                        mass_mg = concentration_mg_per_ml * volume_ml  # mg/mL * mL
                        adjusted_initial_weight_g = mass_mg / 1000  # Convert mg to g
                    elif stock_unit in ['M', 'mM', 'μM', 'uM']:
                        moles = desired_concentration_M * (volume_ml / 1000)  # Convert mL to L
                        adjusted_initial_weight_g = moles * molecular_weight  # in g
                    else:
                        adjusted_initial_weight_g = initial_weight_g  # No change

                    # Update initial weight in component
                    comp['initial_weight'] = adjusted_initial_weight_g / self.conversion_factors[initial_weight_unit]  # Convert back to original unit
                    initial_weight_g = adjusted_initial_weight_g

                else:
                    # Use original initial weight
                    comp['initial_weight'] = initial_weight  # Keep original value

                comp['stock_volume_ml'] = volume_ml
                comp['stock_concentration'] = desired_concentration
                comp['stock_unit'] = comp['stock_unit']
                comp['solvent'] = solvent

                # Calculate cost per mL
                if 'cost' in comp and volume_ml:
                    cost_per_ml = comp['cost'] / volume_ml  # Cost per mL
                    comp['cost_per_ml'] = cost_per_ml
                else:
                    comp['cost_per_ml'] = None

                stock_preparations.append(comp)
        return stock_preparations

    def generate_recipe(self, recipe):
        # First, calculate stock solutions to ensure 'stock_concentration' is set
        self.calculate_stock_solutions()
        output = []
        for item in recipe:
            component_stock = self.find_component_stock(item['name'])
            if component_stock:
                try:
                    volume_ul = self.calculate_volume(component_stock, item)
                    # Calculate cost for the volume added
                    cost = None
                    if 'cost_per_ml' in component_stock and component_stock['cost_per_ml'] is not None:
                        cost = (volume_ul / 1000) * component_stock['cost_per_ml']  # Convert μL to mL
                    component_output = {
                        'name': item['name'],
                        'volume_ul': volume_ul,
                        'cost': cost,
                    }
                    # Add note if working solution is needed
                    if 'working_solution_dilution_factor' in component_stock:
                        dilution_factor = component_stock['working_solution_dilution_factor']
                        component_output['note'] = f"Prepare working solution by diluting the stock {int(dilution_factor)}:1."
                    output.append(component_output)
                except ValueError as e:
                    output.append({
                        'name': item['name'],
                        'volume_ul': None,
                        'cost': None,
                        'note': str(e)
                    })
            else:
                output.append({
                    'name': item['name'],
                    'volume_ul': None,
                    'cost': None,
                    'note': 'Component not found in stock!'
                })
        return output

    def generate_word_document(self, recipe_output, filename='Media_Preparation.docx'):
        document = Document()
        
        # Title
        document.add_heading('Procedure to Prepare Media', 0)

        # Materials Needed
        document.add_heading('Materials Needed:', level=1)
        document.add_paragraph(f"- **Base Media:** {self.base_media['name']}")
        document.add_paragraph(f"- **Serum:** {self.serum['name']}")

        # Stock Solution Preparations
        stock_preparations = self.calculate_stock_solutions()
        if stock_preparations:
            document.add_heading('Stock Solution Preparations:', level=1)
            table = document.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Component'
            hdr_cells[1].text = 'Initial Weight'
            hdr_cells[2].text = 'Stock Concentration'
            hdr_cells[3].text = 'Solvent'
            hdr_cells[4].text = 'Volume to Add'
            hdr_cells[5].text = 'Cost per mL'

            for comp in stock_preparations:
                row_cells = table.add_row().cells
                row_cells[0].text = comp['name']
                initial_weight = comp['initial_weight']
                initial_weight_unit = comp['initial_weight_unit']
                row_cells[1].text = f"{initial_weight:.2f} {initial_weight_unit}"
                stock_concentration = comp['stock_concentration']
                stock_unit = comp['stock_unit']
                row_cells[2].text = f"{stock_concentration:.2f} {stock_unit}"
                row_cells[3].text = comp['solvent']
                volume_ml = comp.get('stock_volume_ml', 'N/A')
                if volume_ml != 'N/A':
                    row_cells[4].text = f"{volume_ml:.2f} mL"
                else:
                    row_cells[4].text = 'N/A'
                cost_per_ml = comp.get('cost_per_ml')
                if cost_per_ml is not None:
                    row_cells[5].text = f"${cost_per_ml:.2f}/mL"
                else:
                    row_cells[5].text = 'N/A'

            document.add_paragraph('Prepare the stock solutions as per the table above.')

        # Media Preparation Steps
        document.add_heading('Media Preparation Steps:', level=1)

        # Calculate the total volume of additives (in μL)
        total_additives_volume_ul = sum(comp['volume_ul'] for comp in recipe_output if comp['volume_ul'] is not None)
        total_additives_volume_ml = total_additives_volume_ul / 1000  # Convert μL to mL

        # Serum volume
        final_volume_ml = self.final_volume_ml
        serum_percentage = self.serum.get('percentage', 10) / 100
        serum_volume_ml = serum_percentage * final_volume_ml

        # Adjust base media volume
        base_media_volume_ml = final_volume_ml - serum_volume_ml - total_additives_volume_ml

        # Ensure base media volume is not negative
        if base_media_volume_ml < 0:
            raise ValueError("Total volume of additives and serum exceeds the final volume. Adjust final volume or component concentrations.")

        # Create table with columns: Step, Component, Desired Concentration, Stock Concentration, Volume to Add (μL), Cost
        table = document.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Step'
        hdr_cells[1].text = 'Component'
        hdr_cells[2].text = 'Desired Concentration'
        hdr_cells[3].text = 'Volume to Add (μL)'
        hdr_cells[4].text = 'Cost'

        step_number = 1

        # Step 1: Add Base Media
        row_cells = table.add_row().cells
        row_cells[0].text = str(step_number)
        row_cells[1].text = self.base_media['name']
        row_cells[2].text = '-'
        row_cells[3].text = f"{base_media_volume_ml * 1000:.2f}"
        row_cells[4].text = '-'
        step_number += 1

        # Step 2: Add Serum
        row_cells = table.add_row().cells
        row_cells[0].text = str(step_number)
        row_cells[1].text = self.serum['name']
        row_cells[2].text = f"{serum_percentage * 100}% v/v"
        row_cells[3].text = f"{serum_volume_ml * 1000:.2f}"
        row_cells[4].text = '-'
        step_number += 1

        # Calculate total cost
        total_cost = 0.0

        # Add Components
        for comp in recipe_output:
            name = comp['name']
            volume_ul = comp['volume_ul']
            cost = comp.get('cost')
            desired_concentration = ''
            stock_concentration = ''
            # Get desired concentration and stock concentration
            for item in self.recipe_data:
                if item['name'] == name:
                    if 'desired_concentration' in item:
                        desired_concentration = f"{item['desired_concentration']} {item['desired_unit']}"
                    elif 'dilution_factor' in item:
                        desired_concentration = f"1:{int(item['dilution_factor'])} dilution"
                    else:
                        desired_concentration = '-'
                    break
            component_stock = self.find_component_stock(name)
            if component_stock:
                stock_concentration = f"{component_stock['stock_concentration']} {component_stock['stock_unit']}"
            else:
                stock_concentration = '-'

            row_cells = table.add_row().cells
            row_cells[0].text = str(step_number)
            row_cells[1].text = name
            row_cells[2].text = desired_concentration
            if volume_ul is not None:
                volume_str = f"{volume_ul:.2f}"
            else:
                volume_str = 'N/A'
            row_cells[3].text = volume_str
            if cost is not None:
                cost_str = f"${cost:.2f}"
                total_cost += cost
            else:
                cost_str = 'N/A'
            row_cells[4].text = cost_str
            step_number += 1

        # Final Mixing Steps
        document.add_paragraph('\n**Final Steps:**')
        document.add_paragraph('- Gently mix all components to ensure thorough mixing.')
        document.add_paragraph('- Avoid creating bubbles.')
        document.add_paragraph('- Use the media immediately or store at 4°C for up to one week.')
        document.add_paragraph('- Protect from light if light-sensitive components are included.')

        # Calculate cost per mL
        cost_per_ml = total_cost / self.final_volume_ml

        # Display Total Cost
        document.add_paragraph(f"\n**Total Cost:**")
        document.add_paragraph(f"- Total Cost of Media: ${total_cost:.2f}")
        document.add_paragraph(f"- Cost per mL of Media: ${cost_per_ml:.2f}/mL")

        # Save the document
        document.save(filename)
        print(f"Word document '{filename}' has been generated successfully.")

    # ... [Include other methods like get_stock_preparation_calculation and get_media_preparation_calculation if necessary]
